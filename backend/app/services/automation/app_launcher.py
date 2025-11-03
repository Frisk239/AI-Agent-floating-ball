from docx import Document
import os
import subprocess
import sys
import time
import pyautogui
import pyperclip

def create_and_open_word_doc(file_name=None):
    """
    创建一个新的空白Word文档并尝试打开它。

    Args:
        file_name (str): 要创建的Word文档的名称。
    """
    try:
        if file_name:
            # 1. 创建一个新的空白文档
            doc = Document()
            # 可以选择添加一些初始内容，这里保持完全空白
            # doc.add_paragraph("这是一个新文档。")

            # 2. 保存文档
            # 保存到桌面
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            full_path = os.path.join(desktop_path, file_name)
            # 确保文件名唯一
            counter = 1
            original_path = full_path
            while os.path.exists(full_path):
                name_without_ext = os.path.splitext(original_path)[0]
                ext = os.path.splitext(original_path)[1]
                new_name = f"{os.path.basename(name_without_ext)}_{counter}{ext}"
                full_path = os.path.join(desktop_path, new_name)
                counter += 1
            doc.save(full_path)
            print(f"文档 '{full_path}' 已创建。")

        else:
            # 1. 创建一个新的空白文档
            doc = Document()
            # 可以选择添加一些初始内容，这里保持完全空白
            # doc.add_paragraph("这是一个新文档。")

            # 2. 保存文档
            # 保存到桌面
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            full_path = os.path.join(desktop_path, "new.docx")
            # 确保文件名唯一
            counter = 1
            original_path = full_path
            while os.path.exists(full_path):
                name_without_ext = os.path.splitext(original_path)[0]
                ext = os.path.splitext(original_path)[1]
                new_name = f"{os.path.basename(name_without_ext)}_{counter}{ext}"
                full_path = os.path.join(desktop_path, new_name)
                counter += 1
            doc.save(full_path)
            print(f"文档 '{full_path}' 已创建。")

        # 4. 根据操作系统打开文件
        # 使用 subprocess.Popen 可以更好地控制新启动的进程
        if sys.platform.startswith('darwin'):  # macOS
            # 使用 'open' 命令
            subprocess.Popen(['open', full_path])
            return "创建并打开文档成功。"
        elif os.name == 'nt':  # Windows
            # 使用 'start' 命令，注意 '/WAIT' 会等待程序关闭，这里不需要
            # 'start' 是 cmd 的内置命令，需要通过 'cmd /c' 调用
            # '/b' 参数可以在后台运行，但不创建新窗口，我们不使用它来确保窗口打开
            subprocess.Popen(['cmd', '/c', 'start', '', full_path], shell=True)
            return "创建并打开文档成功。"
        elif os.name == 'posix':  # Linux
            # 尝试使用 'xdg-open' 命令
            subprocess.Popen(['xdg-open', full_path])
            return "创建并打开文档成功。"
        else:
            print(f"不支持的操作系统: {sys.platform}")
            return "创建或打开文档时出错。"

    except Exception as e:
        print(f"创建或打开文档时出错: {e}")
        return "创建或打开文档时出错。"

def open_netease_music():
    """
    在Windows系统上打开网易云音乐客户端
    该函数会尝试查找并启动网易云音乐程序
    """
    try:
        # 直接使用start命令启动
        subprocess.run(["start", "cloudmusic"], shell=True, check=True)
        print("已尝试启动网易云音乐")
        time.sleep(2)
        return True
    except subprocess.CalledProcessError:
        print("无法通过start命令启动网易云音乐")
        return False


def open_other_apps(app_names: list) -> str:
    """
    通过Windows开始菜单搜索功能打开软件

    当其他函数无法打开某些软件时如QQ，微信等，可以使用此函数通过
    Windows开始菜单的搜索功能来启动应用程序。该函数会模拟按下Win键，
    输入软件名称并回车来启动程序。

    Args:
        app_names (list): 要启动的软件名称列表，例如 ["微信", "QQ音乐", "Photoshop"]

    Returns:
        str: 操作结果描述，包含成功打开的软件列表

    Example:
        >>> open_other_apps(["微信", "QQ音乐"])
        '已打开以下软件: 微信, QQ音乐'
    """
    open_success = []
    for app_name in app_names:
        try:
            # 按下Win键打开开始菜单
            pyautogui.press('win')
            # 等待开始菜单打开
            time.sleep(0.5)
            # 将软件名复制到剪切板
            pyperclip.copy(app_name)
            # 粘贴软件名（使用Ctrl+V）
            pyautogui.hotkey('ctrl', 'v')
            # 等待软件名粘贴完成
            time.sleep(0.5)
            # 按回车键启动软件
            pyautogui.press('enter')
            time.sleep(1)
            open_success.append(app_name)
        except Exception as e:
            return f"打开软件{app_name}时出错: {str(e)}"
    return f"已打开以下软件: {', '.join(open_success)}"


def launch_application_smart(app_name: str) -> dict:
    """
    智能应用启动器 - 双重启动策略

    该函数首先尝试使用预定义的应用映射启动应用程序，
    如果失败则使用搜索启动方式来启动应用程序。

    Args:
        app_name (str): 要启动的应用程序名称

    Returns:
        dict: 包含启动结果的字典
            {
                "success": bool,
                "method": "predefined" | "search",
                "message": str,
                "app_name": str
            }
    """
    # 预定义的应用映射
    app_commands = {
        "notepad": ["notepad.exe"],
        "calculator": ["calc.exe"],
        "explorer": ["explorer.exe"],
        "cmd": ["cmd.exe"],
        "powershell": ["powershell.exe"],
        "chrome": ["start", "chrome"],
        "edge": ["start", "msedge"],
        "vscode": ["code"],
        "word": ["start", "winword"],
        "excel": ["start", "excel"],
        "powerpoint": ["start", "powerpnt"],
        "netease": ["start", "cloudmusic"],  # 网易云音乐
        "wechat": ["start", "WeChat"],  # 微信
        "qq": ["start", "QQ"],  # QQ
        "dingtalk": ["start", "DingTalk"]  # 钉钉
    }

    app_name_lower = app_name.lower()

    # 方法1：尝试预定义启动
    if app_name_lower in app_commands:
        try:
            cmd = app_commands[app_name_lower]
            subprocess.Popen(cmd, shell=True)
            return {
                "success": True,
                "method": "predefined",
                "message": f"应用程序 {app_name} 启动成功",
                "app_name": app_name,
                "command": cmd
            }
        except Exception as e:
            print(f"预定义启动失败: {e}")
            # 预定义启动失败，继续尝试搜索启动

    # 方法2：使用搜索启动
    try:
        # 按下Win键打开开始菜单
        pyautogui.press('win')
        # 等待开始菜单打开
        time.sleep(0.5)
        # 将软件名复制到剪切板
        pyperclip.copy(app_name)
        # 粘贴软件名（使用Ctrl+V）
        pyautogui.hotkey('ctrl', 'v')
        # 等待软件名粘贴完成
        time.sleep(0.5)
        # 按回车键启动软件
        pyautogui.press('enter')
        time.sleep(1)

        return {
            "success": True,
            "method": "search",
            "message": f"应用程序 {app_name} 通过搜索启动成功",
            "app_name": app_name
        }

    except Exception as e:
        return {
            "success": False,
            "method": "failed",
            "message": f"启动应用程序 {app_name} 失败: {str(e)}",
            "app_name": app_name
        }


def control_music_app(actions: list) -> str:
    """
    控制音乐播放应用（如网易云音乐、QQ音乐等）

    该函数通过模拟键盘快捷键来控制音乐播放应用，支持播放控制、音量调节等操作。
    主要适用于网易云音乐，但也兼容其他支持类似快捷键的音乐播放器。

    Args:
        actions (list): 要执行的操作名称列表，支持以下值：
            'play_pause'     → 播放/暂停
            'next_song'      → 下一首
            'previous_song'  → 上一首
            'volume_up'      → 音量增大
            'volume_down'    → 音量减小
            'mini_mode'      → 切换迷你模式
            'like_song'      → 喜欢当前歌曲
            'lyrics_toggle'  → 显示/隐藏歌词
            'mute'           → 静音/取消静音
            'shuffle'        → 切换随机播放
            'repeat'         → 切换重复播放

    Returns:
        str: 操作执行结果描述，包括成功和失败的操作详情

    Example:
        >>> control_music_app(['play_pause', 'next_song'])
        '已成功执行以下操作: play_pause, next_song。'

        >>> control_music_app(['volume_up', 'volume_up'])
        '已成功执行以下操作: volume_up, volume_up。'
    """
    # 根据动作名称，映射到对应的全局快捷键组合
    action_map = {
        'play_pause': ('ctrl', 'alt', 'p'),          # 播放/暂停 (网易云音乐)
        'next_song': ('ctrl', 'alt', 'right'),       # 下一首
        'previous_song': ('ctrl', 'alt', 'left'),    # 上一首
        'volume_up': ('ctrl', 'alt', 'up'),          # 音量增加
        'volume_down': ('ctrl', 'alt', 'down'),      # 音量减少
        'mini_mode': ('ctrl', 'alt', 'o'),           # 迷你模式（网易云默认）
        'like_song': ('ctrl', 'alt', 'l'),           # 喜欢当前歌曲
        'lyrics_toggle': ('ctrl', 'alt', 'd'),       # 歌词开关
        'mute': ('ctrl', 'alt', 'm'),                # 静音
        'shuffle': ('ctrl', 'alt', 's'),             # 随机播放
        'repeat': ('ctrl', 'alt', 'r'),              # 重复播放
    }

    # 分别存储成功和失败的操作
    successful_actions = []
    failed_actions = []

    # 遍历所有请求的操作
    for action in actions:
        # 检查动作是否支持
        if action not in action_map:
            failed_actions.append(f"{action}(不支持的操作)")
            continue

        try:
            # 获取对应快捷键
            keys = action_map[action]
            # 模拟按下快捷键
            pyautogui.hotkey(*keys)

            # 对于音量调节操作，执行两次以获得更明显的效果
            if action in ['volume_up', 'volume_down']:
                time.sleep(0.1)
                pyautogui.hotkey(*keys)

            successful_actions.append(action)

        except Exception as e:
            failed_actions.append(f"{action}(执行时出错: {str(e)})")

    # 等待所有操作执行完成
    if successful_actions:
        time.sleep(len(successful_actions) * 0.1)

    # 构造返回信息
    result_message = ""
    if successful_actions:
        result_message += f"已成功执行以下音乐操作: {', '.join(successful_actions)}。"
    if failed_actions:
        if successful_actions:
            result_message += f" 部分操作执行失败: {', '.join(failed_actions)}。"
        else:
            result_message += f"音乐操作执行失败: {', '.join(failed_actions)}。"

    if not successful_actions and not failed_actions:
        result_message = "没有指定要执行的音乐操作。"

    return result_message


def control_browser_app(actions: list) -> str:
    """
    控制浏览器应用（如Chrome、Edge、Firefox等）

    该函数通过模拟键盘快捷键来控制浏览器，支持标签页管理、书签、历史记录等操作。

    Args:
        actions (list): 要执行的操作名称列表，支持以下值：
            'new_tab'           → 新建标签页
            'close_tab'         → 关闭当前标签页
            'next_tab'          → 切换到下一个标签页
            'previous_tab'      → 切换到上一个标签页
            'reopen_closed_tab' → 重新打开关闭的标签页
            'refresh'           → 刷新页面
            'hard_refresh'      → 强制刷新页面
            'back'              → 后退
            'forward'           → 前进
            'home'              → 回到主页
            'bookmarks'         → 打开书签
            'history'           → 打开历史记录
            'downloads'         → 打开下载
            'fullscreen'        → 全屏/退出全屏
            'zoom_in'           → 放大
            'zoom_out'          → 缩小
            'reset_zoom'        → 重置缩放

    Returns:
        str: 操作执行结果描述，包括成功和失败的操作详情

    Example:
        >>> control_browser_app(['new_tab', 'close_tab'])
        '已成功执行以下浏览器操作: new_tab, close_tab。'

        >>> control_browser_app(['refresh', 'fullscreen'])
        '已成功执行以下浏览器操作: refresh, fullscreen。'
    """
    # 浏览器快捷键映射
    action_map = {
        'new_tab': ('ctrl', 't'),                    # 新建标签页
        'close_tab': ('ctrl', 'w'),                  # 关闭标签页
        'next_tab': ('ctrl', 'tab'),                 # 下一个标签页
        'previous_tab': ('ctrl', 'shift', 'tab'),    # 上一个标签页
        'reopen_closed_tab': ('ctrl', 'shift', 't'), # 重新打开关闭的标签页
        'refresh': ('f5',),                          # 刷新
        'hard_refresh': ('ctrl', 'f5'),              # 强制刷新
        'back': ('alt', 'left'),                     # 后退
        'forward': ('alt', 'right'),                 # 前进
        'home': ('alt', 'home'),                     # 主页
        'bookmarks': ('ctrl', 'shift', 'o'),         # 书签
        'history': ('ctrl', 'h'),                    # 历史记录
        'downloads': ('ctrl', 'j'),                  # 下载
        'fullscreen': ('f11',),                      # 全屏
        'zoom_in': ('ctrl', '+'),                    # 放大
        'zoom_out': ('ctrl', '-'),                   # 缩小
        'reset_zoom': ('ctrl', '0'),                 # 重置缩放
    }

    # 分别存储成功和失败的操作
    successful_actions = []
    failed_actions = []

    # 遍历所有请求的操作
    for action in actions:
        # 检查动作是否支持
        if action not in action_map:
            failed_actions.append(f"{action}(不支持的操作)")
            continue

        try:
            # 获取对应快捷键
            keys = action_map[action]
            # 模拟按下快捷键
            pyautogui.hotkey(*keys)
            successful_actions.append(action)

        except Exception as e:
            failed_actions.append(f"{action}(执行时出错: {str(e)})")

    # 等待所有操作执行完成
    if successful_actions:
        time.sleep(len(successful_actions) * 0.1)

    # 构造返回信息
    result_message = ""
    if successful_actions:
        result_message += f"已成功执行以下浏览器操作: {', '.join(successful_actions)}。"
    if failed_actions:
        if successful_actions:
            result_message += f" 部分操作执行失败: {', '.join(failed_actions)}。"
        else:
            result_message += f"浏览器操作执行失败: {', '.join(failed_actions)}。"

    if not successful_actions and not failed_actions:
        result_message = "没有指定要执行的浏览器操作。"

    return result_message


def control_office_app(actions: list, app_type: str = "word") -> str:
    """
    控制Office应用（如Word、Excel、PowerPoint）

    该函数通过模拟键盘快捷键来控制Office应用程序，支持文档操作、格式设置等。

    Args:
        actions (list): 要执行的操作名称列表
        app_type (str): 应用程序类型，"word"、"excel"、"powerpoint"

    Returns:
        str: 操作执行结果描述
    """
    # Office应用快捷键映射（通用操作）
    common_actions = {
        'save': ('ctrl', 's'),                        # 保存
        'undo': ('ctrl', 'z'),                        # 撤销
        'redo': ('ctrl', 'y'),                        # 重做
        'copy': ('ctrl', 'c'),                        # 复制
        'cut': ('ctrl', 'x'),                         # 剪切
        'paste': ('ctrl', 'v'),                       # 粘贴
        'select_all': ('ctrl', 'a'),                  # 全选
        'find': ('ctrl', 'f'),                        # 查找
        'replace': ('ctrl', 'h'),                     # 替换
        'print': ('ctrl', 'p'),                       # 打印
        'new': ('ctrl', 'n'),                         # 新建
        'open': ('ctrl', 'o'),                        # 打开
        'close': ('ctrl', 'w'),                       # 关闭
    }

    # Word特定操作
    word_actions = {
        'bold': ('ctrl', 'b'),                        # 加粗
        'italic': ('ctrl', 'i'),                      # 斜体
        'underline': ('ctrl', 'u'),                   # 下划线
        'align_left': ('ctrl', 'l'),                  # 左对齐
        'align_center': ('ctrl', 'e'),                # 居中对齐
        'align_right': ('ctrl', 'r'),                 # 右对齐
    }

    # Excel特定操作
    excel_actions = {
        'sum': ('alt', '='),                          # 求和
        'insert_row': ('ctrl', 'shift', '+'),         # 插入行
        'delete_row': ('ctrl', '-'),                  # 删除行
        'insert_column': ('ctrl', 'shift', '+'),      # 插入列
        'delete_column': ('ctrl', '-'),               # 删除列
    }

    # PowerPoint特定操作
    powerpoint_actions = {
        'new_slide': ('ctrl', 'm'),                   # 新幻灯片
        'start_slideshow': ('f5',),                   # 开始放映
        'duplicate_slide': ('ctrl', 'd'),             # 复制幻灯片
    }

    # 合并操作映射
    action_map = common_actions.copy()
    if app_type == "word":
        action_map.update(word_actions)
    elif app_type == "excel":
        action_map.update(excel_actions)
    elif app_type == "powerpoint":
        action_map.update(powerpoint_actions)

    # 分别存储成功和失败的操作
    successful_actions = []
    failed_actions = []

    # 遍历所有请求的操作
    for action in actions:
        # 检查动作是否支持
        if action not in action_map:
            failed_actions.append(f"{action}(不支持的操作)")
            continue

        try:
            # 获取对应快捷键
            keys = action_map[action]
            # 模拟按下快捷键
            pyautogui.hotkey(*keys)
            successful_actions.append(action)

        except Exception as e:
            failed_actions.append(f"{action}(执行时出错: {str(e)})")

    # 等待所有操作执行完成
    if successful_actions:
        time.sleep(len(successful_actions) * 0.1)

    # 构造返回信息
    result_message = ""
    if successful_actions:
        result_message += f"已成功执行以下{app_type.upper()}操作: {', '.join(successful_actions)}。"
    if failed_actions:
        if successful_actions:
            result_message += f" 部分操作执行失败: {', '.join(failed_actions)}。"
        else:
            result_message += f"{app_type.upper()}操作执行失败: {', '.join(failed_actions)}。"

    if not successful_actions and not failed_actions:
        result_message = f"没有指定要执行的{app_type.upper()}操作。"

    return result_message


# --- 主程序 ---
if __name__ == "__main__":
    # 指定文件名
    #document_name = "new_document.docx"
    # 调用函数创建并打开文档
    return_content = create_and_open_word_doc()
    print(return_content)
